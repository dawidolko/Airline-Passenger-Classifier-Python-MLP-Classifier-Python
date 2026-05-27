#!/usr/bin/env python3
"""
Skrypt podmienia TREŚĆ paragrafów w istniejącym docs/dokumentacja_do125148.docx
z projektu Wine Quality na Airline Passenger Satisfaction.
NIE zmienia formatowania, czcionek, stylów, marginesów ani obrazków.
Zachowuje structurę (nagłówki, listy, wcięcia) — podmienia wyłącznie tekst.
"""
from __future__ import annotations

import json
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "docs" / "dokumentacja_do125148.docx"
RESULTS_DIR = ROOT / "results" / "airline"


def zamien_tekst_paragrafu(para: Paragraph, nowy_tekst: str) -> None:
    """
    Podmienia tekst paragrafu zachowując formatowanie pierwszego runa.
    Jeśli paragraf ma wiele runów — kasuje dodatkowe, zostawia pierwszy z nowym tekstem.
    """
    if not para.runs:
        para.text = nowy_tekst
        return

    first_run = para.runs[0]
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)
    first_run.text = nowy_tekst


def podmien_na_indeksie(doc: Document, idx: int, tekst: str) -> None:
    """Podmienia tekst paragrafu o podanym indeksie."""
    if idx < len(doc.paragraphs):
        zamien_tekst_paragrafu(doc.paragraphs[idx], tekst)


def zaktualizuj_docx() -> None:
    comparison = pd.read_csv(RESULTS_DIR / "model_comparison.csv")
    meta = json.loads((RESULTS_DIR / "run_metadata.json").read_text(encoding="utf-8"))
    best = comparison.sort_values("f1_macro", ascending=False).iloc[0]

    doc = Document(DOCX_PATH)

    # --- Nagłówek (para 0) ---
    podmien_na_indeksie(doc, 0, "Sprawozdanie z projektu ze Sztucznej inteligencji:")

    # --- 1. Dane studenta (para 1-6 bez zmian) ---

    # --- 2. Temat projektu (para 6-8) ---
    podmien_na_indeksie(doc, 6, "2. Temat projektu")
    podmien_na_indeksie(
        doc, 7,
        "Klasyfikacja klasy podróży pasażera (Airline Passenger Satisfaction) z wykorzystaniem "
        "sieci neuronowej MLP zaimplementowanej w PyTorch na zbiorze Kaggle. "
        "Projekt porównuje MLP (baseline i tuned po grid search) z Random Forest."
    )
    podmien_na_indeksie(
        doc, 8,
        "Repozytorium: run_experiment.py (eksperyment), airline_project/ (config, model, "
        "preprocessing, experiment), streamlit_app.py (dashboard), results/airline/ (wyniki)."
    )

    # --- 3. Charakterystyka problemu (para 9-13) ---
    podmien_na_indeksie(doc, 9, "3. Charakterystyka problemu")
    podmien_na_indeksie(
        doc, 10,
        "Zadanie to klasyfikacja wieloklasowa (multiclass) nadzorowana. Na podstawie 18 cech "
        "numerycznych (wiek, dystans lotu, 14 ocen usług, opóźnienia) i 3 kategorycznych "
        "(Gender, Customer Type, Type of Travel) model przypisuje pasażera do jednej "
        "z 3 klas podróży: Business, Eco lub Eco Plus."
    )
    podmien_na_indeksie(
        doc, 11,
        "Klasy są niezbalansowane: Business ~48%, Eco ~45%, Eco Plus ~7%. "
        "Model może ignorować mniejszościową klasę i wciąż mieć wysoką accuracy — "
        "stąd konieczność metryki F1-macro i class weights w funkcji straty."
    )
    podmien_na_indeksie(
        doc, 12,
        "Sens praktyczny: predykcja klasy podróży na podstawie profilu i ocen usług "
        "może wspierać personalizację oferty linii lotniczych."
    )
    podmien_na_indeksie(
        doc, 13,
        "Sens edukacyjny: pokazać kompletny pipeline ML z siecią neuronową w PyTorch — "
        "od preprocessingu po grid search, ewaluację i wnioski."
    )

    # --- 4. Liczba instancji (para 14-18) ---
    podmien_na_indeksie(doc, 14, "4. Liczba instancji")
    podmien_na_indeksie(
        doc, 15,
        f"Zbiór Airline Passenger Satisfaction (Kaggle) zawiera {meta['n_records']} rekordów "
        f"(train + test połączone). Po podziale stratyfikowanym 70/15/15:"
    )
    podmien_na_indeksie(doc, 16, f"•  {meta['n_train']} próbek treningowych — fit preprocessora i trening modelu.")
    podmien_na_indeksie(doc, 17, f"•  {meta['n_val']} próbek walidacyjnych — early stopping i scheduler.")
    podmien_na_indeksie(
        doc, 18,
        f"Stratyfikacja zachowuje proporcje klas w każdym podzbiorze. "
        f"Zbiór testowy ({meta['n_test']} próbek) nigdy nie jest widziany podczas treningu."
    )

    # --- 5. Liczba atrybutów (para 19-26) ---
    podmien_na_indeksie(doc, 19, "5. Liczba atrybutów i ich charakterystyka")
    podmien_na_indeksie(
        doc, 20,
        f"Po preprocessingu model otrzymuje {meta['n_features']} cech (18 numerycznych "
        f"po StandardScaler + 6 z OneHotEncoder na 3 kolumnach kategorycznych)."
    )
    podmien_na_indeksie(
        doc, 21,
        "•  Cechy numeryczne (18): Age, Flight Distance, 14 ocen usług (skala 0–5), "
        "Departure Delay in Minutes, Arrival Delay in Minutes."
    )
    podmien_na_indeksie(doc, 22, "•  Cechy kategoryczne (3): Gender, Customer Type, Type of Travel → OneHotEncoder.")
    podmien_na_indeksie(doc, 23, "•  Zmienna docelowa: Class (Business / Eco / Eco Plus) — 3 klasy.")
    podmien_na_indeksie(
        doc, 24,
        "•  Braki danych: jedynie Arrival Delay in Minutes ma ~0.3% NaN; uzupełniane medianą."
    )
    podmien_na_indeksie(
        doc, 25,
        "•  Usunięte kolumny: Unnamed: 0 (indeks CSV), id (identyfikator), "
        "satisfaction (oryginalny binarny target — nie używamy)."
    )
    podmien_na_indeksie(doc, 26, "Konfiguracja kolumn i ścieżek: airline_project/config.py.")

    # --- 6. Preprocessing (para 27-34) ---
    podmien_na_indeksie(doc, 27, "\n6. Preprocessing danych — wstępne przetwarzanie")
    podmien_na_indeksie(
        doc, 28,
        "Preprocessing zaimplementowany w airline_project/preprocessing.py. "
        "Transformacje fitowane wyłącznie na train (brak data leakage)."
    )
    podmien_na_indeksie(doc, 29, "•  Krok 1 — wczytanie CSV (train + test Kaggle), usunięcie kolumn zbędnych.")
    podmien_na_indeksie(
        doc, 30,
        "•  Krok 2 — remove_unrealistic_values: Age poza [0,100], Flight Distance ≤0, "
        "opóźnienia <0 lub >1440 min, oceny poza [0,5] → NaN."
    )
    podmien_na_indeksie(
        doc, 31,
        "•  Krok 3 — SimpleImputer: mediana (numeryczne), moda (kategoryczne) — fit na train."
    )
    podmien_na_indeksie(doc, 32, "•  Krok 4 — OneHotEncoder na Gender, Customer Type, Type of Travel (handle_unknown='ignore').")
    podmien_na_indeksie(doc, 33, "•  Krok 5 — StandardScaler (z-score) na cechach numerycznych — fit na train.")
    podmien_na_indeksie(
        doc, 34,
        "•  Krok 6 — podział stratyfikowany 70/15/15 (random_state=42). "
        "Eksport wyników do results/airline/."
    )

    # --- 7. Projekt modelu (para 35-51) ---
    podmien_na_indeksie(doc, 35, "7. Projekt modelu — MLP (PyTorch), Random Forest")
    podmien_na_indeksie(
        doc, 36,
        "Model główny — AirlineMLP (torch.nn.Module) z konfigurowalnymi warstwami ukrytymi. "
        "Architektura: Linear → BatchNorm1d → ReLU → Dropout (per warstwa ukryta)."
    )
    podmien_na_indeksie(
        doc, 37,
        f"Najlepsza konfiguracja z grid search: {meta['tuned_config']['name']} "
        f"(walidacyjne macro F1 = {meta['best_grid']['best_val_f1_macro']:.4f})."
    )
    podmien_na_indeksie(
        doc, 38,
        "Implementacja: airline_project/model.py (klasa AirlineMLP), "
        "airline_project/experiment.py (trening, grid search, ewaluacja)."
    )
    podmien_na_indeksie(doc, 39, "Konfiguracja MLP:")
    podmien_na_indeksie(doc, 40, "•  Optimizer: AdamW, learning_rate = 1e-3 lub 5e-4, weight_decay = 1e-4.")
    podmien_na_indeksie(doc, 41, "•  Loss: CrossEntropyLoss z class weights (compute_class_weight balanced).")
    podmien_na_indeksie(doc, 42, "•  Scheduler: ReduceLROnPlateau (mode=max, factor=0.5, patience=3).")
    podmien_na_indeksie(doc, 43, "•  Early stopping: patience=5 epok bez poprawy val macro F1.")
    podmien_na_indeksie(doc, 44, "•  Batch size: 1024, max epochs: 30–35.")
    podmien_na_indeksie(doc, 45, "•  Dropout: 0.1 / 0.2 / 0.3 (testowane w grid search).")
    podmien_na_indeksie(doc, 46, "•  Grid search: 3 architektury × 3 dropout × 2 lr = 18 konfiguracji.")
    podmien_na_indeksie(doc, 47, "Model porównawczy:")
    podmien_na_indeksie(doc, 48, "•  RandomForestClassifier (sklearn) z domyślnymi parametrami (100 drzew, bez class weights).")
    podmien_na_indeksie(doc, 49, "•  RF służy jako baseline referencyjny (bez tuningu, bez balansowania).")
    podmien_na_indeksie(doc, 50, "•  Porównanie: MLP baseline vs MLP tuned vs Random Forest na identycznych danych.")
    podmien_na_indeksie(doc, 51, "Ewaluacja: classification_report, confusion matrix, tabela metryk (accuracy, precision/recall/F1 macro).")

    # --- 8. Ocena na zbiorze testowym (para 52-59) ---
    podmien_na_indeksie(doc, 52, "8. Ocena na zbiorze testowym (15%)")
    podmien_na_indeksie(
        doc, 53,
        f"Modele ocenione na {meta['n_test']} próbkach testowych (dane niewidziane w treningu i walidacji)."
    )
    podmien_na_indeksie(doc, 54, "•  Accuracy — odsetek poprawnych klasyfikacji (zdominowane przez duże klasy).")
    podmien_na_indeksie(doc, 55, "•  F1-macro — średnia F1 per klasa z równymi wagami (lepsza przy niezbalansowaniu).")
    podmien_na_indeksie(doc, 56, "•  Precision macro — średnia precyzji per klasa.")
    podmien_na_indeksie(doc, 57, "•  Recall macro — średnia czułości per klasa.")
    podmien_na_indeksie(doc, 58, "•  Macierz pomyłek — analiza pomyłek między Business, Eco, Eco Plus.")

    # --- 9. Analiza rezultatów (para 62-80) ---
    # Find paragraph 62 area
    podmien_na_indeksie(doc, 62, "9. Analiza uzyskanych rezultatów")
    podmien_na_indeksie(doc, 63, "Wyniki na zbiorze testowym (15%):")

    baseline = comparison[comparison["model"] == "MLP baseline"].iloc[0]
    tuned = comparison[comparison["model"] == "MLP tuned"].iloc[0]
    rf = comparison[comparison["model"] == "Random Forest"].iloc[0]

    podmien_na_indeksie(
        doc, 64,
        f"•  MLP baseline (128→64): accuracy={baseline['accuracy']:.4f}, "
        f"F1-macro={baseline['f1_macro']:.4f}"
    )
    podmien_na_indeksie(
        doc, 65,
        f"•  MLP tuned ({meta['tuned_config']['name']}): accuracy={tuned['accuracy']:.4f}, "
        f"F1-macro={tuned['f1_macro']:.4f}"
    )
    podmien_na_indeksie(
        doc, 68,
        f"•  Random Forest (domyślny): accuracy={rf['accuracy']:.4f}, F1-macro={rf['f1_macro']:.4f}"
    )
    podmien_na_indeksie(doc, 69, "Porównanie modeli:")
    podmien_na_indeksie(
        doc, 70,
        "•  MLP tuned osiąga najwyższe F1-macro (~0.653) — najlepsza równowaga między klasami."
    )
    podmien_na_indeksie(
        doc, 71,
        "•  Random Forest ma najwyższą accuracy (~0.864), ale najniższe F1-macro (~0.606) "
        "— ignoruje Eco Plus."
    )
    podmien_na_indeksie(doc, 73, "Interpretacja:")
    podmien_na_indeksie(
        doc, 74,
        "MLP z class weights lepiej radzi sobie z klasą mniejszościową Eco Plus (recall ~41%) "
        "niż Random Forest bez balansowania (recall ~1%). To potwierdza, że accuracy jest "
        "niewystarczającą metryką przy niezbalansowaniu klas."
    )
    podmien_na_indeksie(
        doc, 75,
        "Najtrudniejsza jest klasa Eco Plus (~7%) — w macierzy pomyłek widać, że jest "
        "mylona głównie z Eco (podobne parametry lotu i oceny usług)."
    )
    podmien_na_indeksie(doc, 76, "•  Confusion matrix (Rys.) — Business dobrze rozpoznawany, Eco Plus problematyczny.")
    podmien_na_indeksie(doc, 77, "•  Loss/F1 history (Rys.) — strata spada, val F1 rośnie i stabilizuje się.")
    podmien_na_indeksie(doc, 78, "•  Scheduler LR (Rys.) — learning rate obniżany gdy F1 stagnuje.")
    podmien_na_indeksie(doc, 79, "•  Grid search — większe sieci i wyższy dropout dają lepsze wyniki.")
    podmien_na_indeksie(
        doc, 80,
        "Wniosek: problem wykonalny, MLP z class weights lepiej balansuje klasy niż RF. "
        "Możliwe ulepszenia: RF z class_weight='balanced', SMOTE, Optuna, feature engineering."
    )

    # --- 10. Wnioski (para 81-86) ---
    podmien_na_indeksie(doc, 81, "10. Wnioski")
    podmien_na_indeksie(
        doc, 82,
        f"Zaimplementowano potok ML: dane Airline → preprocessing (czyszczenie, imputacja, "
        f"skalowanie) → MLP w PyTorch (grid search 18 konfiguracji) + Random Forest → "
        f"ewaluacja na teście → wykresy i raporty. "
        f"Najlepszy model: {best['model']} (F1-macro = {best['f1_macro']:.4f})."
    )

    # para 83-86 if they exist
    if len(doc.paragraphs) > 83:
        podmien_na_indeksie(
            doc, 83,
            "Kluczowe obserwacje: (1) class weights w CrossEntropyLoss wymuszają uwagę na Eco Plus; "
            "(2) RF bez balansowania ignoruje klasę mniejszościową; "
            "(3) F1-macro jest rzetelniejszą metryką niż accuracy przy niezbalansowaniu."
        )
    if len(doc.paragraphs) > 84:
        podmien_na_indeksie(
            doc, 84,
            "Ograniczenia: Eco Plus stanowi tylko 7% danych, brak zaawansowanej inżynierii cech, "
            "grid search ograniczony do 18 konfiguracji (kompromis czasowy)."
        )
    if len(doc.paragraphs) > 85:
        podmien_na_indeksie(
            doc, 85,
            "Możliwe ulepszenia: RF z class_weight='balanced', oversampling (SMOTE), "
            "bayesowska optymalizacja (Optuna), feature engineering, ensemble."
        )
    if len(doc.paragraphs) > 86:
        podmien_na_indeksie(
            doc, 86,
            "Projekt spełnia wymagania: MLP jako model główny, RF porównawczy, "
            "preprocessing bez przecieku, metryki F1-macro, macierz pomyłek, wnioski."
        )

    # --- Dodatek A (para 87+) ---
    if len(doc.paragraphs) > 87:
        podmien_na_indeksie(
            doc, 87,
            "Dodatek uzupełnia sprawozdanie o przepływ danych w repozytorium "
            "oraz wykresy wygenerowane z aktualnego uruchomienia."
        )
    if len(doc.paragraphs) > 88:
        podmien_na_indeksie(doc, 88, "A.1. Architektura i przepływ danych")
    if len(doc.paragraphs) > 89:
        podmien_na_indeksie(doc, 89, "Komponenty projektu:")
    if len(doc.paragraphs) > 90:
        podmien_na_indeksie(doc, 90, "•  airline_project/config.py — konfiguracja (ścieżki, cechy, hiperparametry).")
    if len(doc.paragraphs) > 91:
        podmien_na_indeksie(doc, 91, "•  airline_project/preprocessing.py — czyszczenie, podział, imputacja, skalowanie.")
    if len(doc.paragraphs) > 92:
        podmien_na_indeksie(doc, 92, "•  airline_project/model.py — klasa AirlineMLP, wybór urządzenia (CUDA/MPS/CPU).")

    if len(doc.paragraphs) > 93:
        podmien_na_indeksie(doc, 93, "A.2. Eksploracyjna analiza danych (EDA)")
    if len(doc.paragraphs) > 94:
        podmien_na_indeksie(
            doc, 94,
            "EDA: rozkład klas Class (Business ~48%, Eco ~45%, Eco Plus ~7%), "
            "braki danych (~0.3% w Arrival Delay), oceny usług 0–5."
        )
    if len(doc.paragraphs) > 95:
        podmien_na_indeksie(doc, 95, "A.2.1. Rozkład klas (Class)")
    if len(doc.paragraphs) > 97:
        podmien_na_indeksie(
            doc, 97,
            f"Rys. 1. Rozkład klas Class ({meta['n_records']} rekordów). "
            "Business i Eco dominują; Eco Plus (~7%) — klasa mniejszościowa."
        )
    if len(doc.paragraphs) > 98:
        podmien_na_indeksie(doc, 98, "A.2.2. Braki danych")
    if len(doc.paragraphs) > 100:
        podmien_na_indeksie(
            doc, 100,
            "Rys. 2. Braki danych — jedynie Arrival Delay in Minutes ma ~0.3% NaN. "
            "Uzupełniane medianą w pipeline."
        )
    if len(doc.paragraphs) > 101:
        podmien_na_indeksie(doc, 101, "A.2.3. Rozkłady cech numerycznych")
    if len(doc.paragraphs) > 103:
        podmien_na_indeksie(
            doc, 103,
            "Rys. 3. Histogramy cech: Age, Flight Distance, oceny usług. "
            "Różne zakresy — uzasadnia StandardScaler."
        )
    if len(doc.paragraphs) > 104:
        podmien_na_indeksie(doc, 104, "A.2.4. Cechy a klasa podróży")
    if len(doc.paragraphs) > 106:
        podmien_na_indeksie(
            doc, 106,
            "Rys. 4. Związek cech z klasą podróży — Business ma wyższe oceny usług, "
            "Eco Plus trudna do odróżnienia od Eco."
        )
    if len(doc.paragraphs) > 107:
        podmien_na_indeksie(doc, 107, "A.2.5. Macierz pomyłek (confusion matrix)")
    if len(doc.paragraphs) > 109:
        podmien_na_indeksie(
            doc, 109,
            "Rys. 5. Macierze pomyłek side-by-side: Random Forest, MLP baseline, MLP tuned. "
            "RF ignoruje Eco Plus; MLP tuned najlepiej radzi sobie z nią."
        )
    if len(doc.paragraphs) > 110:
        podmien_na_indeksie(doc, 110, "A.2.6. Historia treningu")
    if len(doc.paragraphs) > 112:
        podmien_na_indeksie(
            doc, 112,
            "Rys. 6. Train/val loss i walidacyjne macro F1 po epokach. "
            "Strata spada, F1 rośnie i stabilizuje się przed early stopping."
        )
    if len(doc.paragraphs) > 113:
        podmien_na_indeksie(doc, 113, "A.3. Wyniki")
    if len(doc.paragraphs) > 114:
        podmien_na_indeksie(doc, 114, "A.3.1. Porównanie modeli")
    if len(doc.paragraphs) > 116:
        podmien_na_indeksie(
            doc, 116,
            "Rys. 7. Porównanie modeli na zbiorze testowym. "
            f"Najwyższe F1-macro: {best['model']} ({best['f1_macro']:.4f})."
        )

    # Tabela — podmień zawartość
    table = doc.tables[0]
    headers = ["Model", "Accuracy", "Precision macro", "Recall macro", "F1-macro", "Najlepsza val F1"]
    for i, h in enumerate(headers):
        if i < len(table.rows[0].cells):
            table.rows[0].cells[i].paragraphs[0].runs[0].text = h if table.rows[0].cells[i].paragraphs[0].runs else h

    rows_data = [
        ["MLP baseline", f"{baseline['accuracy']:.4f}", f"{baseline['precision_macro']:.4f}",
         f"{baseline['recall_macro']:.4f}", f"{baseline['f1_macro']:.4f}", "—"],
        ["MLP tuned", f"{tuned['accuracy']:.4f}", f"{tuned['precision_macro']:.4f}",
         f"{tuned['recall_macro']:.4f}", f"{tuned['f1_macro']:.4f}",
         f"{meta['best_grid']['best_val_f1_macro']:.4f}"],
        ["Random Forest", f"{rf['accuracy']:.4f}", f"{rf['precision_macro']:.4f}",
         f"{rf['recall_macro']:.4f}", f"{rf['f1_macro']:.4f}", "—"],
    ]
    for row_idx, row_data in enumerate(rows_data, start=1):
        if row_idx < len(table.rows):
            for col_idx, val in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = val
                else:
                    cell.paragraphs[0].text = val

    # para 117+
    if len(doc.paragraphs) > 117:
        podmien_na_indeksie(doc, 117, "Tabela 1. Wyniki modeli na zbiorze testowym (15%):")
    if len(doc.paragraphs) > 118:
        podmien_na_indeksie(doc, 118, "A.3.2. Macierz pomyłek (MLP tuned, test)")
    if len(doc.paragraphs) > 120:
        podmien_na_indeksie(
            doc, 120,
            f"Rys. 8. Macierz pomyłek MLP tuned na teście ({meta['n_test']} próbek). "
            "Business dobrze rozpoznawany; Eco Plus mylona z Eco."
        )
    if len(doc.paragraphs) > 121:
        podmien_na_indeksie(doc, 121, "A.3.3. Scheduler learning rate")
    if len(doc.paragraphs) > 123:
        podmien_na_indeksie(
            doc, 123,
            "Rys. 9. ReduceLROnPlateau — LR obniżany gdy val macro F1 stagnuje. "
            "Pozwala na dokładniejsze strojenie w końcowych epokach."
        )
    if len(doc.paragraphs) > 124:
        podmien_na_indeksie(doc, 124, "A.3.4. Grid search — najlepsze konfiguracje")
    if len(doc.paragraphs) > 126:
        podmien_na_indeksie(
            doc, 126,
            f"Rys. 10. Top konfiguracje z grid search (18 total). Najlepsza: "
            f"{meta['tuned_config']['name']} (val F1={meta['best_grid']['best_val_f1_macro']:.4f})."
        )
    if len(doc.paragraphs) > 127:
        podmien_na_indeksie(doc, 127, "A.4. Dashboard Streamlit (streamlit_app.py)")
    if len(doc.paragraphs) > 128:
        podmien_na_indeksie(doc, 128, "•  Podsumowanie — metryki najlepszego modelu, tabela porównawcza.")
    if len(doc.paragraphs) > 129:
        podmien_na_indeksie(doc, 129, "•  Raporty — classification_report per model.")
    if len(doc.paragraphs) > 130:
        podmien_na_indeksie(doc, 130, "•  Wykresy — confusion matrices, loss/F1, scheduler LR.")
    if len(doc.paragraphs) > 131:
        podmien_na_indeksie(doc, 131, "•  Wnioski — interpretacja wyników po polsku.")
    if len(doc.paragraphs) > 132:
        podmien_na_indeksie(doc, 132, "•  Pliki — lista artefaktów w results/airline/.")
    if len(doc.paragraphs) > 133:
        podmien_na_indeksie(doc, 133, "•  Przycisk 'Uruchom eksperyment' - trening z poziomu UI.")
    if len(doc.paragraphs) > 134:
        podmien_na_indeksie(doc, 134, "•  Grid search — top 10 konfiguracji z walidacyjnym F1.")
    if len(doc.paragraphs) > 135:
        podmien_na_indeksie(doc, 135, "•  Eksplorator danych — podgląd plików wynikowych.")
    if len(doc.paragraphs) > 136:
        podmien_na_indeksie(
            doc, 136,
            "Dashboard korzysta z plików w results/airline/ wygenerowanych przez "
            "run_experiment.py. Uruchomienie: ./start.sh lub streamlit run streamlit_app.py."
        )

    doc.save(DOCX_PATH)
    print(f"Zaktualizowano treść (bez zmiany formatowania): {DOCX_PATH}")


if __name__ == "__main__":
    zaktualizuj_docx()
