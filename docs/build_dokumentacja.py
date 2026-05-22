#!/usr/bin/env python3
"""
UWAGA: Ten skrypt nadpisuje cały document.xml i usuwa formatowanie szablonu Word.

Do aktualizacji treści i wykresów przy zachowaniu formatu użyj:
    python docs/aktualizuj_dokumentacja.py

Poniższy generator z MD służy tylko do odtworzenia docx od zera (gdy brak szablonu).
"""

from __future__ import annotations

import re
import shutil
import sys
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

MD_PATH = Path(__file__).resolve().parent / "DOKUMENTACJA.md"
OUT_DOCX = Path(__file__).resolve().parent / "dokumentacja_do125148.docx"
TEMPLATE = OUT_DOCX if OUT_DOCX.is_file() else None


def _para(text: str, bold: bool = False) -> str:
    t = escape(text)
    if bold:
        rpr = "<w:rPr><w:b/></w:rPr>"
    else:
        rpr = ""
    return f"<w:p><w:r>{rpr}<w:t xml:space=\"preserve\">{t}</w:t></w:r></w:p>"


def _heading(text: str, level: int) -> str:
    # level 0 -> Heading1, 1 -> Heading2, 2 -> Heading3
    style = {0: "Heading1", 1: "Heading2", 2: "Heading3"}.get(level, "Heading2")
    t = escape(text)
    return (
        f'<w:p><w:pPr><w:pStyle w:val="{style}"/></w:pPr>'
        f"<w:r><w:t xml:space=\"preserve\">{t}</w:t></w:r></w:p>"
    )


def _parse_md_to_body(md: str) -> str:
    parts: list[str] = []
    for line in md.splitlines():
        s = line.strip()
        if not s or s == "---":
            continue
        if s.startswith("# "):
            parts.append(_heading(s[2:], 0))
        elif s.startswith("## "):
            parts.append(_heading(s[3:], 1))
        elif s.startswith("### "):
            parts.append(_heading(s[4:], 2))
        elif s.startswith("```"):
            continue
        elif s.startswith("- "):
            txt = re.sub(r"\*\*([^*]+)\*\*", r"\1", s[2:])
            parts.append(_para("• " + txt))
        elif s.startswith("|"):
            parts.append(_para(s))
        else:
            txt = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
            parts.append(_para(txt))
    parts.append(_para(""))
    return "".join(parts)


def _build_from_template(md_text: str, template: Path, out: Path) -> None:
    body = _parse_md_to_body(md_text)
    sect = (
        '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/>'
        '<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/>'
        "</w:sectPr>"
    )
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}{sect}</w:body></w:document>"
    ).encode("utf-8")

    tmp = out.with_suffix(".tmp.zip")
    with zipfile.ZipFile(template, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = document_xml
            zout.writestr(item, data)
    tmp.replace(out)


def _build_minimal(md_text: str, out: Path) -> None:
    """Minimalny docx bez szablonu — tylko document + content types."""
    body = _parse_md_to_body(md_text)
    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f"<w:body>{body}</w:body></w:document>"
    )
    content_types = """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/document.xml", document_xml.encode("utf-8"))


def main() -> None:
    if not MD_PATH.is_file():
        print(f"Brak {MD_PATH}", file=sys.stderr)
        sys.exit(1)
    md = MD_PATH.read_text(encoding="utf-8")

    try:
        from docx import Document

        doc = Document()
        for line in md.splitlines():
            s = line.strip()
            if not s:
                continue
            if s.startswith("# "):
                doc.add_heading(s[2:], 0)
            elif s.startswith("## "):
                doc.add_heading(s[3:], 1)
            elif s.startswith("### "):
                doc.add_heading(s[4:], 2)
            elif s.startswith("- "):
                doc.add_paragraph(s[2:], style="List Bullet")
            elif s.startswith("```"):
                continue
            else:
                doc.add_paragraph(s)
        doc.save(OUT_DOCX)
        print(f"Zapisano (python-docx): {OUT_DOCX}")
        return
    except ImportError:
        pass

    backup = OUT_DOCX.with_suffix(".docx.bak")
    if OUT_DOCX.is_file() and not backup.is_file():
        shutil.copy2(OUT_DOCX, backup)

    if TEMPLATE and TEMPLATE.is_file():
        _build_from_template(md, TEMPLATE, OUT_DOCX)
        print(f"Zapisano (szablon OOXML): {OUT_DOCX}")
    else:
        _build_minimal(md, OUT_DOCX)
        print(f"Zapisano (minimalny docx): {OUT_DOCX}")


if __name__ == "__main__":
    main()
